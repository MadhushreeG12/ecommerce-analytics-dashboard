"""
Generate a full professional project report (.docx) for:
E-Commerce Customer Analytics Dashboard
Student: Madhu Shree G | USN: 1BM24MC049
Guide: Dr. S. Uma
College: B.M.S. College of Engineering, Bengaluru
Degree: Master of Computer Applications (MCA)
University: VTU, Belgaum
Year: 2025-2026
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page Margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.0)

# ── Styles ────────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_body(doc, text, size=12, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic)
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_numbered(doc, text, size=12):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            hdr_cells[i].width = Inches(col_widths[i])
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F3A5F')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(11)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            row_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_widths:
                row_cells[c_idx].width = Inches(col_widths[c_idx])
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(11)
    return table

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A Project Report on")
set_font(run, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('"Predictive Analytics for Customer Behavior\nin E-Commerce"')
set_font(run, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
run = p.add_run("Submitted in partial fulfilment of requirement for the award of the degree of")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run("MASTER OF COMPUTER APPLICATIONS")
set_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("of")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Visvesvaraya Technological University, Belgaum")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("by")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run("Madhu Shree G")
set_font(run, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("1BM24MC049")
set_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run("Under the guidance of")
set_font(run, size=13, bold=True)

# Guide table
guide_table = doc.add_table(rows=1, cols=2)
guide_table.alignment = WD_TABLE_ALIGNMENT.CENTER
guide_table.style = 'Table Grid'

left_cell = guide_table.rows[0].cells[0]
right_cell = guide_table.rows[0].cells[1]

lp1 = left_cell.add_paragraph()
lp1.add_run("Internal Guide").bold = True
left_cell.add_paragraph("Dr. S. Uma").runs[0].bold = True
left_cell.add_paragraph("Professor")
left_cell.add_paragraph("Department of Computer Applications")
left_cell.add_paragraph("B. M. S. College of Engineering")
left_cell.add_paragraph("Bengaluru – 560019")

rp1 = right_cell.add_paragraph()
rp1.add_run("").bold = True
right_cell.add_paragraph("")
right_cell.add_paragraph("")
right_cell.add_paragraph("")

for cell in guide_table.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(12)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)

# Remove table borders
for row in guide_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        tcPr.append(tcBorders)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("B. M. S. COLLEGE OF ENGINEERING")
set_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("(Autonomous Institution under VTU)")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Bull Temple Road, Basavanagudi, Bengaluru – 560019")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
run = p.add_run("JUNE 2025")
set_font(run, size=13, bold=True)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CERTIFICATE PAGE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "CERTIFICATE", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
doc.add_paragraph()
add_body(doc, "This is to certify that the project report titled Predictive Analytics for Customer Behavior in E-Commerce is a bonafide work carried out by Madhu Shree G (USN: 1BM24MC049) in partial fulfilment of the requirements for the award of the degree of Master of Computer Applications of Visvesvaraya Technological University, Belgaum during the academic year 2025-2026.")
doc.add_paragraph()
add_body(doc, "The project report has been approved as it satisfies the academic requirements in respect of the project work prescribed for the degree.")
doc.add_paragraph()
doc.add_paragraph()

sig_table = doc.add_table(rows=1, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sig_left  = sig_table.rows[0].cells[0]
sig_right = sig_table.rows[0].cells[1]
for para in sig_left.paragraphs:
    para.clear()
for para in sig_right.paragraphs:
    para.clear()

p = sig_left.add_paragraph("Internal Guide")
p.runs[0].bold = True
sig_left.add_paragraph("Dr. S. Uma")
sig_left.add_paragraph("Professor, Dept. of Computer Applications")
sig_left.add_paragraph("B. M. S. College of Engineering")

p = sig_right.add_paragraph("Head of Department")
p.runs[0].bold = True
sig_right.add_paragraph("Department of Computer Applications")
sig_right.add_paragraph("B. M. S. College of Engineering")
sig_right.add_paragraph("Bengaluru – 560019")

for row in sig_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for bn in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{bn}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# DECLARATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "DECLARATION", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
doc.add_paragraph()
add_body(doc, "I, Madhu Shree G (USN: 1BM24MC049), hereby declare that the project report titled Predictive Analytics for Customer Behavior in E-Commerce has been independently carried out by me under the guidance of Dr. S. Uma, Professor, Department of Computer Applications, B. M. S. College of Engineering, Bengaluru.")
doc.add_paragraph()
add_body(doc, "I further declare that this report has not been submitted to any other university or institution for the award of any degree or diploma.")
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Place: Bengaluru")
set_font(run, size=12)
p = doc.add_paragraph()
run = p.add_run("Date:  June 2025")
set_font(run, size=12)
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Madhu Shree G")
set_font(run, size=12, bold=True)
p = doc.add_paragraph()
run = p.add_run("1BM24MC049")
set_font(run, size=12)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "ACKNOWLEDGEMENT", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
doc.add_paragraph()
add_body(doc, "I would like to express my sincere gratitude to all those who made this project possible.")
add_body(doc, "First and foremost, I am deeply grateful to Dr. S. Uma, Professor, Department of Computer Applications, B. M. S. College of Engineering, Bengaluru, for her invaluable guidance, continuous encouragement, and insightful feedback throughout the project.")
add_body(doc, "I extend my heartfelt thanks to the Head of the Department of Computer Applications and the Principal of B. M. S. College of Engineering for providing the necessary infrastructure and resources.")
add_body(doc, "I also thank the faculty members of the Department of Computer Applications for their support and motivation during the course of this project.")
add_body(doc, "Finally, I am grateful to my family and friends for their constant support and encouragement throughout this endeavour.")
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Madhu Shree G")
set_font(run, size=12, bold=True)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "ABSTRACT", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
doc.add_paragraph()
add_body(doc, "Customer churn is one of the most significant challenges faced by e-commerce businesses, directly impacting revenue and long-term sustainability. This project presents the design and development of a comprehensive Predictive Analytics Dashboard for Customer Behavior in E-Commerce, capable of identifying at-risk customers, predicting churn probability, and generating actionable business intelligence insights in real-time.")
add_body(doc, "The system leverages a dataset of 5,630 customers with 20 attributes, encompassing behavioral, transactional, and demographic variables. An extensive Exploratory Data Analysis (EDA) was performed to uncover key churn drivers, including complaint history, customer tenure, satisfaction scores, and engagement patterns. Feature engineering introduced six custom metrics — Customer Lifetime Value (CLV), Recency Score, Engagement Score, Spending Efficiency, High Risk Flag, and Address Diversity Flag — significantly enriching the predictive capability of the models.")
add_body(doc, "Four classification algorithms were trained and evaluated — Logistic Regression, Decision Tree, Random Forest, and Gradient Boosting. The Gradient Boosting model emerged as the best performer with an F1-Score of 91.2% and ROC-AUC of 99.6%, demonstrating excellent precision and recall on the churn detection task. An interactive Streamlit web dashboard was built to visualize EDA findings, compare model performance, display feature importances, and enable a real-time \"What-If\" retention simulator for business users.")
add_body(doc, "The dashboard enables marketing and retention teams to filter high-risk customer segments, export contact lists for CRM integration, and simulate treatment interventions such as complaint resolution, cashback offers, and relationship tenure progression — all updating live churn probability predictions. The result is a production-ready analytics tool that bridges the gap between machine learning models and actionable business decision-making.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "TABLE OF CONTENTS", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
doc.add_paragraph()

toc_entries = [
    ("1.", "Introduction", ""),
    ("1.1", "Purpose", ""),
    ("1.2", "Problem Statement", ""),
    ("1.3", "Motivation", ""),
    ("1.4", "Objective", ""),
    ("1.5", "Scope", ""),
    ("1.6", "Methodology", ""),
    ("2.", "Organization Review", ""),
    ("2.1", "Vision of the Organization", ""),
    ("2.2", "Mission of the Organization", ""),
    ("3.", "Existing and Proposed System", ""),
    ("3.1", "Existing System", ""),
    ("3.1.1", "Limitations of the Existing System", ""),
    ("3.2", "Proposed System", ""),
    ("3.2.1", "Key Features of the Proposed System", ""),
    ("3.2.2", "Solutions Delivered by the Proposed System", ""),
    ("4.", "Software Requirement Specifications (SRS)", ""),
    ("4.1", "Overall Description", ""),
    ("4.2", "Functional Requirements", ""),
    ("4.3", "Non-Functional Requirements", ""),
    ("4.4", "Hardware Requirements", ""),
    ("4.5", "Software Requirements", ""),
    ("4.6", "Assumptions and Dependencies", ""),
    ("4.7", "Constraints", ""),
    ("5.", "Feasibility Study", ""),
    ("5.1", "Technical Feasibility", ""),
    ("5.2", "Operational Feasibility", ""),
    ("5.3", "Economic Feasibility", ""),
    ("6.", "Project Planning", ""),
    ("6.1", "Project Workflow", ""),
    ("6.2", "Gantt Chart", ""),
    ("7.", "System Design", ""),
    ("7.1", "Business Process Flow", ""),
    ("7.2", "Data Flow Diagram (DFD)", ""),
    ("7.3", "Three Tier Architecture", ""),
    ("7.3.1", "Presentation Layer", ""),
    ("7.3.2", "Application Layer (Backend)", ""),
    ("7.3.3", "Data Layer", ""),
    ("8.", "Implementation", ""),
    ("8.1", "Tools and Technologies Used", ""),
    ("9.", "Screenshots", ""),
    ("9.1", "Dashboard Overview Tab", ""),
    ("9.2", "Churn Deep Dive Tab", ""),
    ("9.3", "Behavioral Patterns Tab", ""),
    ("9.4", "Predictive AI Tab", ""),
    ("9.5", "High Risk Customers Tab", ""),
    ("9.6", "Customer Risk Passport", ""),
    ("9.7", "What-If Retention Simulator", ""),
    ("10.", "Test Cases", ""),
    ("11.", "Alignment with Sustainable Development Goals (SDG)", ""),
    ("12.", "Conclusion", ""),
    ("13.", "Team Details", ""),
    ("13.1", "Team Members", ""),
    ("13.2", "Mentor Details", ""),
    ("14.", "Future Enhancements", ""),
    ("15.", "References", ""),
]
for num, title, pg in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    bold = num.endswith('.')
    tab = "" if num.endswith('.') else "    "
    run = p.add_run(f"{tab}{num}  {title}")
    set_font(run, size=11, bold=bold)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. INTRODUCTION", size=15, space_before=0)

add_heading(doc, "1.1 Purpose", size=13)
add_body(doc, "The purpose of this project is to develop a comprehensive, interactive Predictive Analytics Dashboard that enables e-commerce organizations to proactively identify customers at high risk of churning and to take data-driven retention actions. The system consolidates Exploratory Data Analysis (EDA), machine learning-based churn prediction, feature importance visualization, and a live What-If intervention simulator into a single unified web application built using Python and Streamlit.")

add_heading(doc, "1.2 Problem Statement", size=13)
add_body(doc, "E-commerce companies face a critical challenge: customer churn — the phenomenon where customers stop engaging with a platform and switch to competitors. Acquiring a new customer costs five to seven times more than retaining an existing one. Despite the availability of vast amounts of customer behavioral data, most organizations lack the tools to translate this data into timely, actionable predictions. Traditional reporting systems are retrospective, offering insight only after churn has already occurred. There is a pressing need for a proactive, intelligent system that can predict churn in advance, identify its root causes, and recommend personalized retention strategies.")

add_heading(doc, "1.3 Motivation", size=13)
add_body(doc, "Customer retention is a critical success factor in the e-commerce industry. With increasing competition and rising customer acquisition costs, retaining existing customers is far more cost-effective than acquiring new ones. Studies indicate that even a 5% improvement in customer retention can increase profits by 25% to 95%. Despite this, most e-commerce platforms still rely on reactive approaches — reaching out to customers only after they have disengaged. This project is motivated by the need to harness machine learning to build a smarter, proactive approach to customer lifecycle management.")

add_heading(doc, "1.4 Objective", size=13)
add_body(doc, "The key objectives of this project are:")
add_bullet(doc, "To perform Exploratory Data Analysis (EDA) on e-commerce customer data to uncover patterns, correlations, and churn drivers.")
add_bullet(doc, "To engineer meaningful features such as Customer Lifetime Value (CLV), Engagement Score, and High Risk Flag that enhance predictive power.")
add_bullet(doc, "To train and evaluate multiple classification models (Logistic Regression, Decision Tree, Random Forest, and Gradient Boosting) for churn prediction.")
add_bullet(doc, "To build an interactive Streamlit web dashboard that visualizes analytical findings, model performance, and feature importances.")
add_bullet(doc, "To develop a real-time 'What-If' Retention Simulator enabling business users to test intervention strategies and observe live churn risk changes.")
add_bullet(doc, "To segment customers by risk level (Low, Medium, High) and generate targeted marketing campaign lists for CRM integration.")

add_heading(doc, "1.5 Scope", size=13)
add_body(doc, "The scope of this project covers:")
add_bullet(doc, "Analysis of 5,630 e-commerce customer records with 20 attributes including behavioral, transactional, and demographic data.")
add_bullet(doc, "Feature engineering, data cleaning, and preprocessing of the raw dataset.")
add_bullet(doc, "Training and comparison of four classification algorithms for binary churn prediction.")
add_bullet(doc, "Development of a five-tab Streamlit dashboard: Overview, Churn Deep Dive, Behavioral Patterns, Predictive AI, and High Risk Customer Workstation.")
add_bullet(doc, "A live, interactive What-If simulator powered by the trained Gradient Boosting model.")
add_bullet(doc, "Export functionality for filtered customer risk lists in CSV format for CRM integration.")

add_heading(doc, "1.6 Methodology", size=13)
add_body(doc, "This project follows a structured data science lifecycle methodology:")
add_numbered(doc, "Data Collection: Sourced an e-commerce customer dataset in Excel format (.xlsx) from a publicly available repository.")
add_numbered(doc, "Data Cleaning & Preprocessing: Handled missing values using median imputation, corrected inconsistencies, removed duplicates, and encoded categorical features.")
add_numbered(doc, "Exploratory Data Analysis (EDA): Performed statistical analysis, correlation heatmaps, and distribution visualizations using Matplotlib, Seaborn, and Plotly.")
add_numbered(doc, "Feature Engineering: Created six derived features to capture complex customer behavioral patterns.")
add_numbered(doc, "Model Training & Evaluation: Trained four classification models and evaluated them using Accuracy, Precision, Recall, F1-Score, and ROC-AUC.")
add_numbered(doc, "Dashboard Development: Built a multi-tab interactive Streamlit dashboard with live ML inference.")
add_numbered(doc, "Testing & Validation: Conducted functional and performance testing of all dashboard components and model predictions.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ORGANIZATION REVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. ORGANIZATION REVIEW", size=15, space_before=0)
add_body(doc, "B. M. S. College of Engineering (BMSCE) is one of the premier autonomous engineering institutions in India, established in 1946. Affiliated to Visvesvaraya Technological University (VTU), Belgaum, and approved by the All India Council for Technical Education (AICTE), BMSCE has consistently delivered excellence in technical education and research.")

add_heading(doc, "2.1 Vision of the Organization", size=13)
add_body(doc, '"To be a globally acclaimed institution that nurtures innovation, excellence, and ethical leadership through quality education and research."')
add_body(doc, "BMSCE aspires to develop graduates who are equipped to solve real-world engineering and technology challenges while upholding the highest standards of integrity and social responsibility.")

add_heading(doc, "2.2 Mission of the Organization", size=13)
add_body(doc, "The mission of B. M. S. College of Engineering is to:")
add_bullet(doc, "Provide quality education through rigorous academic programs that develop technical competence and innovation.")
add_bullet(doc, "Foster research and entrepreneurship to address industry and societal needs.")
add_bullet(doc, "Build graduates with strong ethical values, leadership skills, and a commitment to lifelong learning.")
add_bullet(doc, "Maintain strong collaboration with industry and academia for holistic development of students.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. EXISTING AND PROPOSED SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. EXISTING AND PROPOSED SYSTEM", size=15, space_before=0)

add_heading(doc, "3.1 Existing System", size=13)
add_body(doc, "Current e-commerce churn management approaches typically rely on:")
add_bullet(doc, "Manual spreadsheet reporting: Marketing teams analyze churn data through Excel sheets, which is time-consuming and prone to human error.")
add_bullet(doc, "Rule-based systems: Simple threshold-based alerts (e.g., 'flag customers inactive for 30 days'), which fail to capture complex multi-variable behavioral patterns.")
add_bullet(doc, "Basic BI tools: Tools like Tableau or Power BI provide retrospective dashboards but lack real-time ML-based churn prediction or intervention simulation capabilities.")
add_bullet(doc, "Siloed data analysis: EDA and modelling are performed separately in different tools, requiring specialized data science expertise to interpret results.")

add_heading(doc, "3.1.1 Limitations of the Existing System", size=12)
add_bullet(doc, "Reactive approach: Churn is identified only after it has occurred, leaving no time for preventive action.")
add_bullet(doc, "No personalized risk scoring: Existing tools treat all disengaged customers equally without computing individual churn probabilities.")
add_bullet(doc, "Lack of intervention simulation: No mechanism exists to test how business actions (e.g., offering a discount, resolving a complaint) would impact churn risk for a specific customer.")
add_bullet(doc, "High dependency on technical staff: Non-technical business users cannot independently explore data or generate risk reports.")
add_bullet(doc, "Poor CRM integration: Export of targeted at-risk customer lists for campaign execution is manual and error-prone.")

add_heading(doc, "3.2 Proposed System", size=13)
add_body(doc, "The proposed system is a full-stack Predictive Analytics Dashboard for E-Commerce Customer Churn, built using Python and Streamlit. It integrates EDA, predictive modeling, and business intelligence into a unified, self-service web application accessible to both data scientists and business users.")

add_heading(doc, "3.2.1 Key Features of the Proposed System", size=12)
add_bullet(doc, "Multi-tab interactive dashboard: Five specialized tabs — Overview, Churn Deep Dive, Behavioral Patterns, Predictive AI, and High Risk Customer Workstation.")
add_bullet(doc, "ML-powered churn prediction: Trained Gradient Boosting classifier (F1=91.2%, AUC=99.6%) scores all customers with live churn probability percentages.")
add_bullet(doc, "Real-time What-If Simulator: Business users can adjust customer attributes (satisfaction score, complaint status, cashback, tenure) and instantly observe how churn probability changes.")
add_bullet(doc, "Dynamic customer segmentation: Automated Low / Medium / High Risk tiers with recommended retention actions (VIP Retention, Re-engagement, Win-back Offer, Discount Coupon).")
add_bullet(doc, "Customer Risk Passport: Full risk profile for any individual customer including a gauge chart, engagement metrics, and CLV.")
add_bullet(doc, "CSV export: Filtered risk-segmented campaign lists can be exported directly to CSV for CRM integration.")
add_bullet(doc, "Sidebar filters: Dynamic filters for Gender, Marital Status, and City Tier update all charts and metrics in real-time.")

add_heading(doc, "3.2.2 Solutions Delivered by the Proposed System", size=12)
add_body(doc, "The proposed system delivers the following solutions:")
add_table(doc,
    ["Problem", "Solution Delivered"],
    [
        ["No proactive churn detection", "Real-time churn probability scoring for all 5,630 customers using Gradient Boosting."],
        ["No individual risk profiling", "Customer Risk Passport with gauge chart, profile details, and tailored recommendations."],
        ["No intervention simulation", "Live What-If Simulator: adjust 4 treatment parameters and see churn risk update instantly."],
        ["No self-service analytics", "Interactive Streamlit dashboard with no-code exploration tools for business users."],
        ["No automated segmentation", "Automated 3-tier risk segmentation with 4 recommendation categories."],
        ["No CRM export", "One-click CSV export of filtered campaign contact lists."],
    ],
    col_widths=[2.5, 3.5]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. SRS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. SOFTWARE REQUIREMENT SPECIFICATIONS (SRS)", size=15, space_before=0)

add_heading(doc, "4.1 Overall Description", size=13)
add_body(doc, "The system is a web-based analytics application built using Python 3.x and the Streamlit framework. It operates on structured e-commerce customer data and uses pre-trained Scikit-learn machine learning models to generate real-time predictions. The dashboard is hosted locally and can be accessed via a standard web browser. All data processing, feature engineering, model inference, and visualization occur in-memory within the Python runtime environment.")

add_heading(doc, "4.2 Functional Requirements", size=13)
add_table(doc,
    ["FR No.", "Requirement Description"],
    [
        ["FR-01", "The system shall load and preprocess the E-Commerce Dataset from Excel or CSV format on startup."],
        ["FR-02", "The system shall provide dynamic sidebar filters for Gender, Marital Status, and City Tier."],
        ["FR-03", "The system shall display five key performance metric cards: Churn Rate, Avg Tenure, Avg Satisfaction, Complaint Rate, and Avg Cashback."],
        ["FR-04", "The Overview Tab shall display churn distribution pie chart, tenure histogram, marital status bar chart, and satisfaction box plot."],
        ["FR-05", "The Churn Deep Dive Tab shall display a correlation heatmap, churn-by-complaint bar chart, and churn-by-city-tier chart."],
        ["FR-06", "The Behavioral Patterns Tab shall display login device sunburst chart, payment mode bar chart, and order-vs-cashback scatter plot."],
        ["FR-07", "The Predictive AI Tab shall display a grouped bar chart comparing all model metrics and a styled classification performance table."],
        ["FR-08", "The Predictive AI Tab shall display the feature importance chart from the trained Gradient Boosting model."],
        ["FR-09", "The system shall load the saved Gradient Boosting model (best_model_gb.pkl) and compute churn probabilities for all customers."],
        ["FR-10", "The High Risk Tab shall provide segment explorer with multi-select risk tier and action filters plus a search-by-customer-ID field."],
        ["FR-11", "The system shall display a risk-color-coded, sortable customer table with churn probability, risk level, and recommended action."],
        ["FR-12", "The system shall provide a one-click CSV export of the filtered customer list."],
        ["FR-13", "The Customer Risk Passport shall display an individual gauge chart and full profile details for any selected customer."],
        ["FR-14", "The What-If Simulator shall accept four user-adjustable inputs and return updated churn probability in real-time."],
    ],
    col_widths=[0.8, 5.5]
)

add_heading(doc, "4.3 Non-Functional Requirements", size=13)
add_table(doc,
    ["NFR No.", "Category", "Description"],
    [
        ["NFR-01", "Performance", "Dashboard shall load all tabs within 5 seconds on a standard laptop with 8GB RAM."],
        ["NFR-02", "Reliability", "The ML model shall produce consistent predictions across repeated runs on the same input data."],
        ["NFR-03", "Usability", "The dashboard shall be operable by non-technical business users with no Python knowledge."],
        ["NFR-04", "Scalability", "The system shall handle up to 50,000 customer records without significant performance degradation."],
        ["NFR-05", "Maintainability", "All code shall be modular, commented, and organized for easy extension or model replacement."],
        ["NFR-06", "Security", "The system shall not expose raw customer data beyond the local machine environment."],
    ],
    col_widths=[0.8, 1.3, 4.2]
)

add_heading(doc, "4.4 Hardware Requirements", size=13)
add_table(doc,
    ["Component", "Minimum Requirement"],
    [
        ["Processor", "Intel Core i5 / AMD Ryzen 5 or equivalent (2.0 GHz+)"],
        ["RAM", "8 GB (16 GB recommended for large datasets)"],
        ["Storage", "At least 2 GB free disk space for dataset, models, and dependencies"],
        ["Display", "1366 × 768 resolution or higher"],
        ["Network", "Not required (fully offline capable)"],
    ],
    col_widths=[2.0, 4.3]
)

add_heading(doc, "4.5 Software Requirements", size=13)
add_table(doc,
    ["Software / Library", "Version", "Purpose"],
    [
        ["Python",              "3.10+",  "Core programming language"],
        ["Streamlit",          "1.35+",  "Web dashboard framework"],
        ["Pandas",             "2.0+",   "Data manipulation and analysis"],
        ["NumPy",              "1.24+",  "Numerical operations"],
        ["Scikit-learn",       "1.3+",   "Machine learning models and preprocessing"],
        ["Plotly",             "5.18+",  "Interactive data visualizations"],
        ["Matplotlib",         "3.7+",   "Static chart generation"],
        ["Seaborn",            "0.12+",  "Statistical visualization"],
        ["Joblib",             "1.3+",   "Model serialization and loading"],
        ["openpyxl",           "3.1+",   "Excel dataset loading"],
        ["OS / pathlib",       "Built-in","Cross-platform file path handling"],
    ],
    col_widths=[1.8, 1.0, 3.5]
)

add_heading(doc, "4.6 Assumptions and Dependencies", size=13)
add_bullet(doc, "The E-Commerce Dataset (E Commerce Dataset.xlsx) is available in the archive (3)/ subdirectory of the project folder.")
add_bullet(doc, "The pre-trained model (best_model_gb.pkl) and scaler (scaler.pkl) are generated by running phase2_ml.py before launching the dashboard.")
add_bullet(doc, "The dashboard is assumed to run on a single local machine; multi-user deployment or authentication is outside the current scope.")
add_bullet(doc, "Python 3.10+ and all listed dependencies in requirements.txt are installed in the environment.")

add_heading(doc, "4.7 Constraints", size=13)
add_bullet(doc, "The system is designed for offline, local deployment. Cloud hosting or authentication mechanisms are not implemented.")
add_bullet(doc, "The predictive model is trained on a specific dataset; accuracy may vary if applied to significantly different e-commerce domains without retraining.")
add_bullet(doc, "The feature importance chart (fig8_feature_importance.png) is pre-generated and static; it does not update dynamically with sidebar filters.")
add_bullet(doc, "CSV export generates a flat file and does not directly push to CRM systems.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. FEASIBILITY STUDY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. FEASIBILITY STUDY", size=15, space_before=0)

add_heading(doc, "5.1 Technical Feasibility", size=13)
add_body(doc, "The project is technically feasible using widely available, open-source technologies. Python, Streamlit, Scikit-learn, Pandas, and Plotly are all mature, well-documented libraries with extensive community support. The Gradient Boosting classifier achieves excellent performance (F1: 91.2%, AUC: 99.6%) on the dataset, confirming the technical viability of the ML-based prediction approach. The Streamlit framework enables rapid development of a production-quality, browser-based interactive dashboard without requiring front-end development expertise. All system requirements — data loading, preprocessing, model inference, visualization, and export — are achievable using the selected technology stack.")

add_heading(doc, "5.2 Operational Feasibility", size=13)
add_body(doc, "The system is designed for non-technical business users. The Streamlit-based interface requires no programming knowledge and can be accessed through a standard web browser after a one-time setup. Sidebar filters, multi-select dropdowns, sliders, and search fields make data exploration intuitive. The What-If Simulator allows retention teams to instantly test intervention strategies without requiring a data scientist's involvement. The CSV export feature integrates seamlessly with existing CRM workflows. Training of marketing and retention staff on the dashboard is expected to require less than two hours.")

add_heading(doc, "5.3 Economic Feasibility", size=13)
add_body(doc, "The project is economically viable as it relies entirely on open-source, zero-cost libraries and frameworks. There are no licensing fees for Python, Streamlit, Scikit-learn, Pandas, Plotly, or any other component used. The only infrastructure requirement is a standard personal computer or laptop. The business value generated — through reduction in customer churn, improved retention campaign targeting, and proactive intervention — significantly outweighs the minimal development and operational cost.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. PROJECT PLANNING
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. PROJECT PLANNING", size=15, space_before=0)

add_heading(doc, "6.1 Project Workflow", size=13)
add_body(doc, "The project was executed in two structured phases:")
add_body(doc, "Phase 1 - EDA & Data Preparation:")
add_numbered(doc, "Dataset acquisition and initial exploration.")
add_numbered(doc, "Data cleaning: Missing value imputation, duplicate removal, and inconsistency correction.")
add_numbered(doc, "Exploratory Data Analysis: Distribution plots, correlation heatmaps, churn-by-segment analysis.")
add_numbered(doc, "Feature engineering: Creation of CLV, Recency Score, Engagement Score, Spending Efficiency, High Risk Flag, and Address Diversity Flag.")
add_numbered(doc, "Dataset preprocessing: Label encoding and Standard Scaling.")

add_body(doc, "Phase 2 - Predictive Modeling & Dashboard Development:")
add_numbered(doc, "Training of four classification models and three regression models.")
add_numbered(doc, "Model evaluation using accuracy, precision, recall, F1-Score, ROC-AUC, and cross-validation F1.")
add_numbered(doc, "Saving best model (Gradient Boosting) and scaler as .pkl files.")
add_numbered(doc, "Development of the five-tab Streamlit dashboard.")
add_numbered(doc, "Integration of the What-If Simulator with live model inference.")
add_numbered(doc, "Testing, debugging, and final deployment.")

add_heading(doc, "6.2 Gantt Chart", size=13)
add_table(doc,
    ["Phase / Task", "Week 1", "Week 2", "Week 3", "Week 4", "Week 5", "Week 6", "Week 7", "Week 8"],
    [
        ["Dataset Acquisition & EDA",         "✓", "✓", "",  "",  "",  "",  "",  ""],
        ["Data Cleaning & Preprocessing",     "",  "✓", "✓", "",  "",  "",  "",  ""],
        ["Feature Engineering",               "",  "",  "✓", "✓", "",  "",  "",  ""],
        ["Model Training & Evaluation",       "",  "",  "",  "✓", "✓", "",  "",  ""],
        ["Dashboard Development (Tabs 1-3)",  "",  "",  "",  "",  "✓", "✓", "",  ""],
        ["Dashboard Development (Tabs 4-5)",  "",  "",  "",  "",  "",  "✓", "✓", ""],
        ["What-If Simulator & Testing",       "",  "",  "",  "",  "",  "",  "✓", "✓"],
        ["Documentation & Report",            "",  "",  "",  "",  "",  "",  "✓", "✓"],
    ],
    col_widths=[2.2, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55, 0.55]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. SYSTEM DESIGN", size=15, space_before=0)

add_heading(doc, "7.1 Business Process Flow", size=13)
add_body(doc, "The end-to-end data and business process flow of the system is as follows:")
add_numbered(doc, "Raw Data Ingestion: The E-Commerce Dataset is loaded from an Excel file. Missing values are imputed, inconsistencies corrected, and duplicates removed.")
add_numbered(doc, "Feature Engineering: Six engineered features (CLV, RecencyScore, EngagementScore, SpendingEfficiency, HighRisk, AddressDiversityFlag) are computed from raw attributes.")
add_numbered(doc, "Encoding & Scaling: Categorical variables are Label Encoded; numerical features are Standard Scaled using the saved scaler.")
add_numbered(doc, "Model Inference: The pre-trained Gradient Boosting model scores all customers with a churn probability (0–100%).")
add_numbered(doc, "Risk Segmentation: Customers are categorized into Low Risk (≤30%), Medium Risk (31–70%), or High Risk (>70%) tiers.")
add_numbered(doc, "Recommendation Engine: Each customer is assigned a recommended retention action based on CLV, Engagement Score, and Recency Score.")
add_numbered(doc, "Dashboard Presentation: All insights are rendered in the five-tab Streamlit dashboard with interactive filters and real-time chart updates.")
add_numbered(doc, "What-If Simulation: Business users adjust attributes; the system re-computes churn probability using the live model and presents the delta.")
add_numbered(doc, "Export: Filtered customer lists are exported as CSV files for CRM campaign execution.")

add_heading(doc, "7.2 Data Flow Diagram (DFD)", size=13)
add_body(doc, "Level 0 — Context Diagram:")
add_bullet(doc, "External Entity: Business User → System: Dashboard Application → Output: Insights, Risk Scores, Campaign Lists")

add_body(doc, "Level 1 — Process Breakdown:")
add_bullet(doc, "Process 1.0: Data Loading & Cleaning — Input: Raw Excel file; Output: Cleaned DataFrame.")
add_bullet(doc, "Process 2.0: Feature Engineering — Input: Cleaned DataFrame; Output: Enriched DataFrame (26 features).")
add_bullet(doc, "Process 3.0: Model Inference — Input: Encoded & Scaled features; Output: Churn probabilities.")
add_bullet(doc, "Process 4.0: Dashboard Rendering — Input: Enriched + scored DataFrame; Output: Interactive visualizations.")
add_bullet(doc, "Process 5.0: What-If Simulation — Input: User-modified attributes; Output: Updated churn probability delta.")
add_bullet(doc, "Process 6.0: Export — Input: Filtered risk list; Output: campaign_list.csv.")

add_heading(doc, "7.3 Three Tier Architecture", size=13)

add_heading(doc, "7.3.1 Presentation Layer", size=12)
add_body(doc, "The Presentation Layer is the Streamlit-powered web interface rendered in a browser. It consists of:")
add_bullet(doc, "Five interactive tabs with Plotly charts, styled Pandas DataFrames, and custom HTML/CSS components.")
add_bullet(doc, "A sidebar with dynamic multi-select filters for Gender, Marital Status, and City Tier.")
add_bullet(doc, "A Customer Risk Passport with a Plotly Gauge chart and live What-If sliders/checkboxes.")
add_bullet(doc, "Key Metrics cards at the top displaying Churn Rate, Avg Tenure, Avg Satisfaction, Complaint Rate, and Avg Cashback.")

add_heading(doc, "7.3.2 Application Layer (Backend)", size=12)
add_body(doc, "The Application Layer handles all business logic and ML inference in Python:")
add_bullet(doc, "Data loading, cleaning, and feature engineering pipeline (dashboard.py, phase2_ml.py).")
add_bullet(doc, "Label Encoding, Standard Scaling, and feature selection consistent with the trained model's expectations.")
add_bullet(doc, "Gradient Boosting model inference (joblib.load + predict_proba).")
add_bullet(doc, "Risk tier classification and recommendation engine logic.")
add_bullet(doc, "What-If simulation engine: re-encodes mutated customer attributes and invokes model inference in real-time.")

add_heading(doc, "7.3.3 Data Layer", size=12)
add_body(doc, "The Data Layer manages all persistent data assets:")
add_bullet(doc, "E Commerce Dataset.xlsx: Raw input data (5,630 rows × 20 columns).")
add_bullet(doc, "cleaned_ecommerce_dataset.csv: Pre-processed dataset for rapid fallback loading.")
add_bullet(doc, "best_model_gb.pkl: Serialized Gradient Boosting classifier.")
add_bullet(doc, "scaler.pkl: Serialized StandardScaler fitted on training data.")
add_bullet(doc, "model_summary.json: Dictionary of all model performance metrics for dashboard visualization.")
add_bullet(doc, "fig8_feature_importance.png: Pre-generated static feature importance chart.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. IMPLEMENTATION", size=15, space_before=0)

add_heading(doc, "8.1 Tools and Technologies Used", size=13)
add_table(doc,
    ["Tool / Technology", "Category", "Usage"],
    [
        ["Python 3.10+",       "Language",       "Primary development language for data processing and ML."],
        ["Streamlit 1.35+",    "Web Framework",  "Interactive multi-tab dashboard with live ML inference."],
        ["Scikit-learn",       "ML Library",     "Model training (Logistic Regression, DT, RF, GB), StandardScaler, LabelEncoder."],
        ["Pandas",             "Data Processing","DataFrame operations, EDA, feature engineering."],
        ["NumPy",              "Computation",    "Numerical array operations and mathematical functions."],
        ["Plotly Express",     "Visualization",  "Interactive pie charts, bar charts, scatter plots, heatmaps, sunburst."],
        ["Plotly Go (Gauge)",  "Visualization",  "Customer Risk Gauge chart in the Risk Passport."],
        ["Matplotlib/Seaborn", "Visualization",  "Static chart generation in phase2_ml.py (saved as .png files)."],
        ["Joblib",             "Serialization",  "Saving and loading trained model and scaler as .pkl files."],
        ["openpyxl",           "Data I/O",       "Reading the Excel dataset (.xlsx)."],
        ["JSON",               "Data Storage",   "Storing model performance metrics in model_summary.json."],
        ["VS Code",            "IDE",            "Primary development environment."],
    ],
    col_widths=[1.6, 1.3, 3.4]
)

add_body(doc, "The ML pipeline was implemented as follows:")
add_body(doc, "1. Data Preparation (phase2_ml.py):")
add_bullet(doc, "Dataset loaded, cleaned, and feature-engineered (6 new features).")
add_bullet(doc, "Label encoded and Standard Scaled for model compatibility.")
add_bullet(doc, "80/20 train-test split with stratification to preserve churn class balance.")

add_body(doc, "2. Model Training & Evaluation:")
add_body(doc, "Four classifiers were trained and evaluated on the test set:")
add_table(doc,
    ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC", "CV F1"],
    [
        ["Gradient Boosting",  "97.16%", "95.40%", "87.37%", "91.21%", "99.61%", "93.55%"],
        ["Random Forest",      "95.29%", "86.24%", "85.79%", "86.02%", "98.41%", "87.51%"],
        ["Decision Tree",      "80.02%", "45.10%", "84.74%", "58.87%", "87.98%", "65.69%"],
        ["Logistic Regression","79.13%", "43.59%", "80.53%", "56.56%", "86.83%", "58.40%"],
    ],
    col_widths=[1.5, 0.9, 0.9, 0.7, 0.9, 0.9, 0.8]
)
add_body(doc, "Gradient Boosting outperformed all other models with the highest F1-Score (91.21%) and ROC-AUC (99.61%), and was selected as the production model saved as best_model_gb.pkl.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. SCREENSHOTS", size=15, space_before=0)

screenshots = [
    ("9.1 Dashboard Overview Tab",
     "The Overview Tab provides a comprehensive summary of the customer base. It features five key metric cards (Churn Rate, Average Tenure, Average Satisfaction, Complaint Rate, and Average Cashback) at the top. Below, a donut-style pie chart shows the Churn Distribution (Active vs. Churned customers), a histogram displays the Tenure Distribution, a bar chart shows the Marital Status Distribution, and a box plot visualizes the Satisfaction Score Spread. Sidebar filters for Gender, Marital Status, and City Tier update all charts dynamically."),
    ("9.2 Churn Deep Dive Tab",
     "The Churn Deep Dive Tab reveals the key drivers of customer churn through analytical visualizations. A colour-coded Correlation Heatmap highlights the strength and direction of relationships between numeric features and the Churn variable — Complain and Tenure emerge as the strongest predictors. Two bar charts display the Churn Rate by Complaint Status and Churn Rate by City Tier, enabling marketers to identify the highest-risk customer segments at a glance."),
    ("9.3 Behavioral Patterns Tab",
     "The Behavioral Patterns Tab explores customer engagement and transactional behavior. A Sunburst Chart shows the preferred login device breakdown by Gender, revealing that mobile users form the majority. A bar chart ranks Preferred Payment Modes by frequency. A scatter plot overlaid with OLS trend lines maps Order Count against Cashback Amount, colour-coded by Churn status, exposing behavioral differences between active and churned customers."),
    ("9.4 Predictive AI Tab",
     "The Predictive AI Tab presents the machine learning results side-by-side. On the left, a grouped Plotly bar chart compares all five performance metrics (Accuracy, Precision, Recall, F1-Score, ROC-AUC) across all four classifiers for direct visual comparison. On the right, a styled Pandas DataFrame table presents the same metrics with dark header styling and green highlights on the best-performing metric for each column. Below, the Feature Importance chart from the Gradient Boosting model reveals the Top 15 most influential features, with Complain, Tenure, and CLV ranking highest."),
    ("9.5 High Risk Customer Workstation Tab",
     "The High Risk Customer Workstation provides a comprehensive, interactive environment for retention teams. The Segment Explorer section includes multi-select filters for Risk Level and Recommended Action, a Customer ID search field, and four summary metrics cards. The resulting filtered table displays customers sorted by descending churn probability, colour-coded by Risk Level (red for High Risk, yellow for Medium, green for Low), with the Recommended Action and key behavioral metrics."),
    ("9.6 Customer Risk Passport",
     "The Customer Risk Passport enables business users to select any individual customer from the full dataset and view their complete risk profile. A Plotly Gauge Chart vividly displays the customer's churn probability on a colour-coded dial (green = low risk, amber = medium, red = high). Below the gauge, a detailed profile card shows Gender, Marital Status, City Tier, Tenure, Satisfaction Score, Complaint Status, Order Count, Cashback, Engagement Score, and CLV."),
    ("9.7 What-If Retention Simulator",
     "The What-If Retention Simulator empowers marketing and retention teams to test intervention strategies in real-time without any technical expertise. Users can adjust four parameters — Complaint Resolution, Satisfaction Score, Cashback Reward Boost (₹), and Relationship Progression (+ months) — using checkboxes and sliders. The system immediately re-runs the Gradient Boosting model on the modified customer record and displays the New Churn Probability and New Risk Level, along with a success message showing the exact reduction in churn risk achieved."),
]

for title, description in screenshots:
    add_heading(doc, title, size=13)
    add_body(doc, description)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Screenshot Placeholder]")
    set_font(run, size=11, italic=True, color=(150, 150, 150))
    doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 10. TEST CASES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. TEST CASES", size=15, space_before=0)
add_body(doc, "The following test cases were executed to validate the functional correctness and reliability of the dashboard:")
add_table(doc,
    ["TC No.", "Test Description", "Input", "Expected Output", "Result"],
    [
        ["TC-01", "Dataset loading from Excel", "E Commerce Dataset.xlsx", "5,630 rows loaded, 0 missing values after cleaning", "PASS"],
        ["TC-02", "Sidebar filter — Gender", "Select 'Male' only", "All charts update to show only Male customers", "PASS"],
        ["TC-03", "Sidebar filter — City Tier", "Select Tier 1 only", "Metrics recalculate for Tier 1 customers only", "PASS"],
        ["TC-04", "Model loading", "best_model_gb.pkl present", "Model loads without error; predictions generated for all rows", "PASS"],
        ["TC-05", "Churn probability generation", "5,630 customer records", "All customers scored with probability in [0, 100]", "PASS"],
        ["TC-06", "Risk tier segmentation", "Churn prob = 85%", "Customer classified as 🔴 High Risk", "PASS"],
        ["TC-07", "Risk tier segmentation", "Churn prob = 45%", "Customer classified as 🟡 Medium Risk", "PASS"],
        ["TC-08", "Risk tier segmentation", "Churn prob = 15%", "Customer classified as 🟢 Low Risk", "PASS"],
        ["TC-09", "Customer ID search", "Input '50005' in search box", "Table filters to show only customer 50005", "PASS"],
        ["TC-10", "CSV export", "Click 'Export Filtered List'", "CSV downloaded with correct columns and data", "PASS"],
        ["TC-11", "What-If Simulator — resolve complaint", "Complain=1 → resolved", "Churn probability decreases", "PASS"],
        ["TC-12", "What-If Simulator — cashback boost", "Add ₹100 cashback", "Churn probability updates reflecting new CLV", "PASS"],
        ["TC-13", "What-If Simulator — tenure increase", "Add 12 months tenure", "Churn probability decreases significantly", "PASS"],
        ["TC-14", "Feature importance chart", "fig8_feature_importance.png exists", "Image displayed with correct caption (Gradient Boosting)", "PASS"],
        ["TC-15", "Fallback data loading", "Excel file missing; CSV present", "System loads cleaned_ecommerce_dataset.csv successfully", "PASS"],
    ],
    col_widths=[0.5, 1.8, 1.4, 1.8, 0.7]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 11. SDG ALIGNMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "11. ALIGNMENT WITH SUSTAINABLE DEVELOPMENT GOALS (SDG)", size=15, space_before=0)
add_body(doc, "This project aligns with the United Nations Sustainable Development Goals (SDGs) in the following ways:")
add_table(doc,
    ["SDG Goal", "Alignment with Project"],
    [
        ["SDG 8: Decent Work and Economic Growth",
         "By enabling e-commerce companies to retain customers more efficiently, the system reduces revenue loss and supports sustainable business growth — contributing to economic prosperity and the creation of stable employment within the e-commerce ecosystem."],
        ["SDG 9: Industry, Innovation and Infrastructure",
         "The project demonstrates applied innovation by combining machine learning, feature engineering, and interactive data visualization to solve a real-world business problem. It exemplifies the responsible use of technology and AI for industrial value creation."],
        ["SDG 12: Responsible Consumption and Production",
         "By identifying disengaged customers early and targeting them with personalized offers, the system reduces wasteful blanket marketing spend and promotes more efficient, targeted resource allocation in business operations."],
        ["SDG 17: Partnerships for the Goals",
         "The system's CSV export and CRM integration features enable collaboration between data science teams and marketing departments, fostering inter-departmental partnership through data-driven decision-making."],
    ],
    col_widths=[2.0, 4.3]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 12. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "12. CONCLUSION", size=15, space_before=0)
add_body(doc, "This project successfully developed a comprehensive, production-ready Predictive Analytics Dashboard for E-Commerce Customer Churn. By integrating exploratory data analysis, advanced feature engineering, and state-of-the-art machine learning with an interactive Streamlit web interface, the system bridges the gap between raw data and actionable business intelligence.")
add_body(doc, "The Gradient Boosting model emerged as the best classifier with an outstanding F1-Score of 91.21% and ROC-AUC of 99.61%, demonstrating the power of ensemble learning for churn prediction on complex, multi-dimensional customer data. The six engineered features — particularly Customer Lifetime Value (CLV), Engagement Score, and the High Risk Flag — significantly enhanced the predictive capability and business interpretability of the models.")
add_body(doc, "The five-tab Streamlit dashboard provides marketing and retention teams with an intuitive, self-service platform to explore customer risk profiles, generate targeted campaign lists, and simulate the impact of retention strategies in real-time — all without requiring programming expertise. The What-If Retention Simulator, powered by live Gradient Boosting inference, enables data-driven intervention decisions at the individual customer level.")
add_body(doc, "In summary, this project demonstrates the practical application of data science and machine learning in solving a critical e-commerce business challenge, delivering measurable value through early churn identification, personalized retention recommendations, and seamless CRM integration capabilities.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 13. TEAM DETAILS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "13. TEAM DETAILS", size=15, space_before=0)

add_heading(doc, "13.1 Team Members", size=13)
add_table(doc,
    ["Sl. No.", "Name", "USN", "Department", "College"],
    [
        ["1", "Madhu Shree G", "1BM24MC049", "MCA", "B. M. S. College of Engineering, Bengaluru"],
    ],
    col_widths=[0.5, 1.5, 1.3, 0.8, 2.5]
)

add_heading(doc, "13.2 Mentor Details", size=13)
add_table(doc,
    ["Role", "Name", "Designation", "Department", "Institution"],
    [
        ["Internal Guide", "Dr. S. Uma", "Professor", "Dept. of Computer Applications", "B. M. S. College of Engineering, Bengaluru"],
    ],
    col_widths=[1.1, 1.2, 1.1, 1.7, 1.5]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 14. FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "14. FUTURE ENHANCEMENTS", size=15, space_before=0)
add_body(doc, "The following enhancements are planned for future versions of the system:")
add_numbered(doc, "Cloud Deployment: Host the Streamlit dashboard on a cloud platform (e.g., AWS, GCP, or Streamlit Community Cloud) to enable multi-user access with authentication and role-based permissions.")
add_numbered(doc, "Real-Time Data Pipeline: Replace batch processing with a real-time streaming pipeline using Apache Kafka or AWS Kinesis, enabling live churn scoring as new customer events occur.")
add_numbered(doc, "Deep Learning Models: Explore LSTM-based sequential models or TabNet for improved prediction accuracy on temporal customer behavioral patterns.")
add_numbered(doc, "Automated Model Retraining: Implement MLOps pipelines using MLflow or Kubeflow to automatically retrain and redeploy models as new data becomes available.")
add_numbered(doc, "Direct CRM Integration: Build API connectors to popular CRM systems (e.g., Salesforce, HubSpot) to push retention recommendations and trigger automated campaigns directly from the dashboard.")
add_numbered(doc, "Multi-Language Support: Add localization and internationalization support to extend the dashboard to global e-commerce operations.")
add_numbered(doc, "Explainability (XAI): Integrate SHAP (SHapley Additive exPlanations) values to provide model-level and prediction-level explanations for individual churn decisions, enhancing trust and regulatory compliance.")
add_numbered(doc, "Mobile-Responsive Design: Optimize the Streamlit layout for mobile and tablet access to support field sales and retention teams on-the-go.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 15. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "15. REFERENCES", size=15, space_before=0)
refs = [
    "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785–794). ACM.",
    "Scikit-learn Documentation (2024). Gradient Boosting Classifier. Retrieved from https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.GradientBoostingClassifier.html",
    "Streamlit Documentation (2024). Build a Machine Learning App. Retrieved from https://docs.streamlit.io",
    "Pandas Development Team (2024). Pandas: Powerful Data Structures for Data Analysis. Retrieved from https://pandas.pydata.org/docs/",
    "Plotly Technologies Inc. (2024). Plotly Python Open Source Graphing Library. Retrieved from https://plotly.com/python/",
    "Baesens, B., Verbeke, W., & Viaene, S. (2014). New Insights into Churn Prediction in the Telecommunication Sector: A Profit Driven Data Mining Approach. European Journal of Operational Research, 218(1), 211–229.",
    "Hadden, J., Tiwari, A., Roy, R., & Ruta, D. (2007). Computer Assisted Customer Churn Management: State-of-the-art and Future Trends. Computers & Operations Research, 34(10), 2902–2917.",
    "E-Commerce Dataset. Retrieved from Kaggle: https://www.kaggle.com/datasets/ankitverma2010/ecommerce-customer-churn-analysis-and-prediction",
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
    "VanderPlas, J. (2016). Python Data Science Handbook: Essential Tools for Working with Data. O'Reilly Media.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f"[{i}]  {ref}")
    set_font(run, size=11)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("Project_Report_Madhu_Shree_G.docx")
print("Report saved as: Project_Report_Madhu_Shree_G.docx")
